from __future__ import annotations

import json
import os
from pathlib import Path
import subprocess
import sys
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from lxml import etree

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": WORD_NS, "m": MATH_NS}
TOOL_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = Path(__file__).resolve().parents[4]


def call(root: Path, payload: dict[str, object]) -> dict[str, object]:
    env = os.environ.copy()
    env["TIANCE_WORKSPACE_ROOT"] = str(root)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONPATH"] = os.pathsep.join(
        [str(TOOL_ROOT / "program"), str(REPOSITORY_ROOT / "1_PythonServer")]
    )
    completed = subprocess.run(
        [sys.executable, str(TOOL_ROOT / "program" / "main.py")],
        input=json.dumps(payload, ensure_ascii=False),
        capture_output=True,
        text=True,
        encoding="utf-8",
        env=env,
        check=False,
    )
    return json.loads(completed.stdout)


def document_xml(path: Path) -> etree._Element:
    with ZipFile(path) as package:
        return etree.fromstring(package.read("word/document.xml"))


def test_create_uses_content_aware_table_widths_and_native_formula(tmp_path: Path) -> None:
    output = tmp_path / "result.docx"
    result = call(
        tmp_path,
        {
            "action": "create",
            "output_path": output.name,
            "elements": [
                {
                    "type": "table",
                    "rows": [
                        ["编号", "非常长的工作内容说明"],
                        ["1", "这一列需要明显更多宽度，以减少无意义换行。"],
                    ],
                },
                {"type": "equation", "latex": r"\sum_{i=1}^{N} x_i"},
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath("count(.//w:tblLayout[@w:type='fixed'])", namespaces=NS) == 1
    widths = [int(value) for value in root.xpath(".//w:tblGrid[1]/w:gridCol/@w:w", namespaces=NS)]
    assert len(widths) == 2
    assert widths[1] > widths[0]
    assert root.xpath("count(.//m:oMath)", namespaces=NS) >= 1


def test_invalid_formula_is_preserved_as_text_with_warning(tmp_path: Path) -> None:
    output = tmp_path / "invalid.docx"
    result = call(
        tmp_path,
        {
            "action": "create",
            "output_path": output.name,
            "elements": [{"type": "equation", "latex": r"\frac{a{b}"}],
        },
    )
    assert result["ok"] is True
    assert any("花括号未配对" in warning for warning in result["warnings"])
    root = document_xml(output)
    assert "\\frac{a{b}" in "".join(root.itertext())


def test_selection_replaces_text_and_removes_inline_equation_without_touching_anchors(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "edited.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("左边界")
    equation = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    equation.append(math_run)
    paragraph._p.append(equation)
    paragraph.add_run("待替换内容")
    paragraph.add_run("右边界")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "replace",
                    "content_mode": "text",
                    "content": "新内容",
                    "style": {"bold": True, "color": "C00000"},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert "".join(root.xpath(".//w:t/text()", namespaces=NS)) == "左边界新内容右边界"
    assert root.xpath("count(.//m:oMath)", namespaces=NS) == 0
    assert root.xpath("count(.//w:r[w:t='新内容']/w:rPr/w:b)", namespaces=NS) == 1
    assert root.xpath(".//w:r[w:t='新内容']/w:rPr/w:color/@w:val", namespaces=NS) == ["C00000"]


def test_zero_width_selection_only_allows_insert(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    inserted = tmp_path / "inserted.docx"
    doc = Document()
    doc.add_paragraph("锚点后文")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": inserted.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "锚点"},
                    "action": "insert",
                    "content": "插入",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "".join(document_xml(inserted).xpath(".//w:t/text()", namespaces=NS)) == "锚点插入后文"

    rejected = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": "rejected.docx",
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "锚点"},
                    "action": "delete",
                }
            ],
        },
    )
    assert rejected["ok"] is False
    assert "零长度选区只能执行 insert" in rejected["error"]


def test_selection_formats_only_existing_selected_text(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("开始")
    target_run = paragraph.add_run("需要格式化")
    target_run.font.size = Pt(10)
    target_run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    paragraph.add_run("结束")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "开始", "end_anchor": "结束"},
                    "action": "format",
                    "style": {"bold": True, "italic": True, "font_size": 14, "color": "1F4E79"},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath("count(.//w:r[w:t='需要格式化']/w:rPr/w:b)", namespaces=NS) == 1
    assert root.xpath("count(.//w:r[w:t='需要格式化']/w:rPr/w:i)", namespaces=NS) == 1
    assert root.xpath(".//w:r[w:t='需要格式化']/w:rPr/w:color/@w:val", namespaces=NS) == ["1F4E79"]
    assert root.xpath("count(.//w:r[w:t='开始']/w:rPr/w:b)", namespaces=NS) == 0
    assert root.xpath("count(.//w:r[w:t='结束']/w:rPr/w:b)", namespaces=NS) == 0


def test_format_patch_does_not_reset_unspecified_existing_style(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    doc = Document()
    paragraph = doc.add_paragraph("左")
    run = paragraph.add_run("目标")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    paragraph.add_run("右")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "右"},
                    "action": "format",
                    "style": {"bold": True},
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    assert root.xpath(".//w:r[w:t='目标']/w:rPr/w:color/@w:val", namespaces=NS) == ["123456"]
    assert root.xpath(".//w:r[w:t='目标']/w:rPr/w:sz/@w:val", namespaces=NS) == ["30"]


def test_cross_block_markdown_replace_removes_old_table_and_writes_native_content(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "markdown.docx"
    doc = Document()
    doc.add_paragraph("左边界")
    doc.add_paragraph("旧段落")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "旧表格"
    doc.add_paragraph("右边界")
    doc.save(source)

    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "replace",
                    "content_mode": "markdown",
                    "content": "## 新章节\n\n带有行内公式 $x^2+y^2=z^2$。\n\n| 项目 | 公式 |\n| --- | --- |\n| 面积 | $S=\\pi r^2$ |",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    root = document_xml(output)
    text = "".join(root.xpath(".//w:t/text()", namespaces=NS))
    assert "左边界" in text and "右边界" in text
    assert "旧段落" not in text and "旧表格" not in text
    assert "新章节" in text and "面积" in text
    assert root.xpath("count(.//w:tbl)", namespaces=NS) == 1
    assert root.xpath("count(.//m:oMath)", namespaces=NS) >= 2
    body_paragraphs = [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        for paragraph in root.xpath("/w:document/w:body/w:p", namespaces=NS)
    ]
    assert body_paragraphs == ["左边界", "新章节", "带有行内公式 。", "右边界"]


def test_selection_extract_and_dry_run_report_range_without_writing(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "unused.docx"
    doc = Document()
    doc.add_paragraph("左边界需要提取右边界")
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "dry_run": True,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左边界", "end_anchor": "右边界"},
                    "action": "extract",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert result["data"]["operations"][0]["selection"]["selected_text"] == "需要提取"
    assert not output.exists()

    preview = call(
        tmp_path,
        {
            "action": "inspect",
            "input_path": source.name,
            "inspect": {"selection": {"start_anchor": "左边界", "end_anchor": "右边界"}},
        },
    )
    assert preview["ok"] is True, preview
    assert preview["data"]["selection"]["selected_text"] == "需要提取"


def test_same_table_cell_selection_is_supported_but_cross_cell_range_is_rejected(tmp_path: Path) -> None:
    source = tmp_path / "table.docx"
    output = tmp_path / "table_edited.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "左选区右"
    table.cell(0, 1).text = "另一单元格"
    doc.save(source)
    result = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": output.name,
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "右"},
                    "action": "replace",
                    "content": "新",
                }
            ],
        },
    )
    assert result["ok"] is True, result
    assert "左新右" in "".join(document_xml(output).xpath(".//w:t/text()", namespaces=NS))

    rejected = call(
        tmp_path,
        {
            "action": "edit",
            "input_path": source.name,
            "output_path": "bad.docx",
            "operations": [
                {
                    "type": "selection",
                    "selection": {"start_anchor": "左", "end_anchor": "另一单元格"},
                    "action": "delete",
                }
            ],
        },
    )
    assert rejected["ok"] is False
    assert "不同表格单元格" in rejected["error"]
