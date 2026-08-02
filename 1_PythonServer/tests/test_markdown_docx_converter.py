from __future__ import annotations

import base64
import os
import subprocess
import time
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.services.document_conversion import MarkdownDocxService
from app.services.document_conversion.markdown_docx import (
    browser_capture,
    browser_renderer,
    docx_package,
)
from app.services.document_conversion.markdown_docx.converter import convert_markdown_to_docx
from app.services.document_conversion.markdown_docx.formula_converter import (
    OMML,
    latex_to_omml,
    preprocess_latex,
    validate_latex,
)
from app.services.document_conversion.markdown_docx.latex_extensions import (
    normalize_for_image,
    normalize_for_omml,
)


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z3jcAAAAASUVORK5CYII="
)


def test_markdown_docx_converter_handles_table_edge_cases(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        """## 表格边界

| 指标 | 完成情况 | 评定 |
|------|----------|:--:|
| 完成标准 | 谭总明确：<80分不合格，80-90合格底线，>90完成较好 | — |

| 仅一列 |
|:------:|
| 单列内容 |
| 第二行 |

普通文本：A | B。
""",
    )

    assert len(document.tables) == 2
    assert _table_rows(document.tables[0]) == [
        ["指标", "完成情况", "评定"],
        ["完成标准", "谭总明确：<80分不合格，80-90合格底线，>90完成较好", "—"],
    ]
    assert _table_rows(document.tables[1]) == [
        ["仅一列"],
        ["单列内容"],
        ["第二行"],
    ]
    assert "普通文本：A | B。" in [paragraph.text for paragraph in document.paragraphs]


def test_table_keeps_code_literal_and_converts_unquoted_formula(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        r"""| 代码 | 公式 |
| --- | --- |
| `$a\mid b$` | $a\mid b$ |
""",
    )

    assert document.tables[0].cell(1, 0).text == "$a\\mid b$"
    assert len(document.tables[0].cell(1, 1)._element.xpath(".//m:oMath")) == 1


def test_markdown_docx_converter_normalizes_visible_double_quotes(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        r"""中文"内容"，English "content"。

"**跨格式**"；**"粗体"**；["链接"](https://example.com/?q="raw")

行内代码 `config="raw"`，尺寸 15"，显式 \"直引号\"，孤立 "引号。

| 名称 |
| --- |
| "表格内容" |

```json
{"key": "value"}
```
""",
    )

    assert document.paragraphs[0].text == "中文“内容”，English “content”。"
    assert document.paragraphs[1].text == "“跨格式”；“粗体”；“链接”"
    assert document.paragraphs[2].text == (
        '行内代码 config="raw"，尺寸 15"，显式 "直引号"，孤立 "引号。'
    )
    assert document.paragraphs[3].text == '{"key": "value"}'
    assert document.tables[0].cell(1, 0).text == "“表格内容”"
    assert any(
        relationship.target_ref == 'https://example.com/?q="raw"'
        for relationship in document.part.rels.values()
    )


def test_markdown_docx_converter_aligns_list_markers_with_body_indent(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        """正文段落用于比较缩进位置。

- 一级项目
  - 二级项目
1. 编号项目
""",
    )

    body, bullet, nested_bullet, ordered = document.paragraphs
    assert round(body.paragraph_format.first_line_indent.inches, 2) == 0.28
    assert bullet._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val") == ["0"]
    assert nested_bullet._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val") == ["1"]
    assert ordered._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val") == ["0"]
    assert bullet.text == "一级项目"
    assert nested_bullet.text == "二级项目"
    assert ordered.text == "编号项目"


def test_latex_preprocessing_respects_command_boundaries():
    normalized = preprocess_latex(r"\int_0^\infty x\in A \cdot y \cdots")

    assert normalized == r"\int_0^∞ x∈ A ⋅ y ⋯"
    assert "∈t" not in normalized
    assert "∈fty" not in normalized
    assert "⋅s" not in normalized


def test_latex_preprocessing_normalizes_unsupported_named_operators():
    normalized = preprocess_latex(
        r"\csch x+\sech y+\arg\max f(x)+\arg\min g(y)+\argmax h(x)+\argmin k(y)"
    )

    assert normalized == (
        r"\operatorname{csch} x+\operatorname{sech} y+"
        r"\operatorname{argmax} f(x)+\operatorname{argmin} g(y)+"
        r"\operatorname{argmax} h(x)+\operatorname{argmin} k(y)"
    )
    omml, error = latex_to_omml(normalized)
    assert error == ""
    assert omml is not None


def test_markdown_docx_converter_separates_currency_code_and_inline_math(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        """商品价格为 $100，折扣后为 $80；真正的公式是 $p=0.8p_0$。

代码 `` `$x+y$` `` 应作为代码；外部 $x-y$ 应作为公式。

转义金额 \\$100 不应成为公式。

短公式：$x$、$xy$、$x1$、$123$。
""",
    )

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "$100" in paragraphs[0]
    assert "$80" in paragraphs[0]
    assert "`$x+y$`" in paragraphs[1]
    assert "$100" in paragraphs[2]
    assert len(document._element.xpath(".//m:oMath")) == 6


def test_markdown_docx_converter_supports_bracket_display_math(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        r"""\[
\int_0^1x^2\,dx=\frac13
\]
""",
    )

    assert len(document._element.xpath(".//m:oMath")) == 1


def test_markdown_docx_converter_supports_multiline_parenthesized_math(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        r"""\(
\sum_{i=1}^{n} i=\frac{n(n+1)}{2}
\)

正文中的 \(x+1\) 仍然是行内公式。
""",
    )

    assert len(document._element.xpath(".//m:oMath")) == 2
    assert any("正文中的" in paragraph.text for paragraph in document.paragraphs)


def test_markdown_docx_converter_degrades_suspicious_formulas_without_losing_tail(
    tmp_path: Path,
):
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        r"""$$\frac{a+b}{c+d$$

$$\thisCommandDefinitelyDoesNotExist{a+b}+c$$

$$\color{red}{x^2}+\textcolor{blue}{y^2}=\colorbox{yellow}{z^2}$$

尾部哨兵 $x+1$。
""",
    )

    assert len(warnings) == 3
    assert any("花括号未配对" in warning for warning in warnings)
    assert any("未识别命令" in warning for warning in warnings)
    assert any("colorbox 已转换为 Word 可编辑底纹" in warning for warning in warnings)
    assert not any("已作为图片插入" in warning for warning in warnings)
    assert len(document._element.xpath(".//m:oMath")) == 2
    assert len(document.inline_shapes) == 0
    assert any("尾部哨兵" in paragraph.text for paragraph in document.paragraphs)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert r"\frac{a+b}{c+d" in document_text
    assert r"\thisCommandDefinitelyDoesNotExist{a+b}+c" in document_text
    assert r"\color{red}{x^2}" not in document_text


def test_unknown_declaration_color_warns_without_losing_editable_formula(tmp_path: Path):
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        r"$$C05:\quad {\color{brandPrimary}x+y}$$",
    )

    assert len(document._element.xpath(".//m:oMath")) == 1
    assert any("brandPrimary" in warning for warning in warnings)


def test_valid_advanced_formulas_convert_without_empty_word_placeholders():
    d05, d05_error = latex_to_omml(
        preprocess_latex(r"D05:\quad \sideset{_{a}^{b}}{_{c}^{d}}\sum_{n=1}^{\infty}x_n")
    )
    assert d05_error == ""
    assert d05 is not None
    nary_operands = [element.find(f"{OMML}e") for element in d05.iter(f"{OMML}nary")]
    assert all("".join(operand.itertext()) for operand in nary_operands if operand is not None)
    assert "x" in "".join(d05.itertext())

    for source in (
        r"F05:\quad {}^{14}_{6}\mathrm C+{}_{Z}^{A}X+\prescript{a}{b}{T}^{c}_{d}",
        r"I03:\quad R^{\rho}{}_{\sigma\mu\nu}",
    ):
        normalized, notices = normalize_for_omml(source)
        assert notices == ()
        omml, error = latex_to_omml(preprocess_latex(normalized))
        assert error == ""
        assert omml is not None
        for tag in ("sSub", "sSup", "sSubSup"):
            for script in omml.iter(f"{OMML}{tag}"):
                operand = script.find(f"{OMML}e")
                assert operand is not None
                assert len(operand) > 0 or bool((operand.text or "").strip())


def test_aligned_nary_repair_does_not_nest_matrix_operands():
    source = r"""\begin{aligned}
\frac{\partial \mathcal{L}}{\partial w} &= \frac{1}{n}\sum_{i=1}^{n} \frac{\partial \ell_i}{\partial w} \\
\sqrt{a^2 + b^2} &= \sqrt{c^2 + d^2} \\
\sum_{k=1}^{\infty} \frac{1}{k^2} &= \frac{\pi^2}{6}
\end{aligned}"""
    normalized, notices = normalize_for_omml(source)

    omml, error = latex_to_omml(preprocess_latex(normalized))

    assert notices == ()
    assert error == ""
    assert omml is not None
    nested_operands = [
        child
        for operand in omml.iter(f"{OMML}e")
        for child in operand
        if child.tag == f"{OMML}e"
    ]
    nary_texts = [
        "".join(operand.itertext())
        for nary in omml.iter(f"{OMML}nary")
        if (operand := nary.find(f"{OMML}e")) is not None
    ]
    assert nested_operands == []
    assert nary_texts == ["∂ℓi∂w", "1k2"]


def test_genfrac_is_expanded_to_editable_word_formula():
    normalized, notices = normalize_for_omml(
        r"G02:\quad \genfrac{\langle}{\rangle}{1pt}{}{a+b}{c+d}"
        r"+\genfrac{(}{)}{0pt}{1}{n}{k}"
    )

    assert r"\genfrac" not in normalized
    assert notices == (
        "genfrac 的自定义分数线宽已按 Word 标准分数线显示",
        "genfrac 的字号样式由 Word 公式自动适配",
    )
    omml, error = latex_to_omml(preprocess_latex(normalized))
    assert error == ""
    assert omml is not None
    assert len(list(omml.iter(f"{OMML}f"))) == 1
    assert len(list(omml.iter(f"{OMML}eqArr"))) == 1


def test_style_extensions_are_normalized_to_editable_word_formulas():
    colored, color_notices = normalize_for_omml(
        r"J02:\quad \color{red}{x^2}+\textcolor{blue}{y^2}=\colorbox{yellow}{$z^2$}"
    )
    assert color_notices == (
        "colorbox 已转换为 Word 可编辑底纹；复杂公式的背景会按公式节点分段显示",
    )
    assert r"\color" not in colored
    assert r"\colorbox" not in colored
    assert r"\style{color:#FF0000}{x^2}" in colored
    assert r"\style{color:#0000FF}{y^2}" in colored
    assert r"\style{background:FFFF00}{z^2}" in colored
    colored_omml, color_error = latex_to_omml(preprocess_latex(colored))
    assert color_error == ""
    assert colored_omml is not None
    assert [node.get(qn("w:val")) for node in colored_omml.iter(qn("w:color"))] == [
        "FF0000",
        "FF0000",
        "0000FF",
        "0000FF",
    ]
    assert [node.get(qn("w:fill")) for node in colored_omml.iter(qn("w:shd"))] == [
        "FFFF00",
        "FFFF00",
    ]

    cases = (
        (
            r"J03:\quad \cancel{x}+\bcancel{y}+\xcancel{z}+\cancelto{0}{a-a}",
            (r"\cancel", r"\bcancel", r"\xcancel", r"\cancelto"),
            ("删除线样式已忽略", "cancelto 的删除箭头已忽略"),
        ),
        (
            r"J04:\quad \boxed{E=mc^2}+\fbox{$a+b$}+\bbox[5px,border:1px solid red]{x+y}",
            (r"\fbox", r"\bbox", "$"),
            ("fbox 已按 Word 原生公式框显示", "bbox 的自定义边距和边框样式已按 Word 原生公式框显示"),
        ),
    )

    for source, removed_fragments, expected_notices in cases:
        normalized, notices = normalize_for_omml(source)
        assert all(fragment not in normalized for fragment in removed_fragments)
        assert all(any(expected in notice for notice in notices) for expected in expected_notices)
        omml, error = latex_to_omml(preprocess_latex(normalized))
        assert error == ""
        assert omml is not None

    boxed, _ = normalize_for_omml(cases[1][0])
    assert boxed.count(r"\boxed") == 3
    cancelto, _ = normalize_for_omml(cases[0][0])
    assert r"\overset{0}{a-a}" in cancelto


def test_formula_colors_support_scopes_nested_structures_and_safe_fallbacks():
    normalized, notices = normalize_for_omml(
        r"{\color{blue}\sqrt{x+1}}+\textcolor{red}{\frac{a+b}{c}}"
        r"+\colorbox{#0f0}{$z^2$}"
    )
    assert notices == (
        "colorbox 已转换为 Word 可编辑底纹；复杂公式的背景会按公式节点分段显示",
    )
    omml, error = latex_to_omml(preprocess_latex(normalized))
    assert error == ""
    assert omml is not None
    colors = [node.get(qn("w:val")) for node in omml.iter(qn("w:color"))]
    assert "0000FF" in colors
    assert colors.count("FF0000") == 3
    radical_properties = next(omml.iter(f"{OMML}radPr"))
    fraction_properties = next(omml.iter(f"{OMML}fPr"))
    assert next(radical_properties.iter(qn("w:color"))).get(qn("w:val")) == "0000FF"
    assert next(fraction_properties.iter(qn("w:color"))).get(qn("w:val")) == "FF0000"
    assert [node.get(qn("w:fill")) for node in omml.iter(qn("w:shd"))] == [
        "00FF00",
        "00FF00",
    ]

    unsupported, unsupported_notices = normalize_for_omml(r"\textcolor{brandPrimary}{x+y}")
    assert unsupported == "x+y"
    assert unsupported_notices == ("暂不支持颜色 brandPrimary，已保留数学内容并忽略颜色",)


def test_extension_packages_and_commutative_diagrams_render_as_images(tmp_path: Path, monkeypatch):
    rendered: list[str] = []
    _mock_formula_renderer(monkeypatch, rendered)
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        r"""$$J02:\quad \color{red}{x^2}+\textcolor{blue}{y^2}=\colorbox{yellow}{$z^2$}$$

$$J03:\quad \cancel{x}+\bcancel{y}+\xcancel{z}+\cancelto{0}{a-a}$$

$$J04:\quad \boxed{E=mc^2}+\fbox{$a+b$}+\bbox[5px,border:1px solid red]{x+y}$$

$$J05:\quad \ce{2H2 + O2 -> 2H2O}+\pu{9.81 m s^-2}$$

$$J06:\quad \xymatrix{A\ar[r]^f\ar[d]_g&B\ar[d]^h\\C\ar[r]_k&D}$$

$$
\begin{tikzcd}
J07:\quad A \arrow[r,"f"] \arrow[d,"g"'] & B \arrow[d,"h"] \\
C \arrow[r,"k"'] & D
\end{tikzcd}
$$
""",
    )

    assert len(document._element.xpath(".//m:oMath")) == 3
    assert len(document.inline_shapes) == 3
    image_warnings = [warning for warning in warnings if "已作为图片插入" in warning]
    assert len(image_warnings) == 3
    assert not any("颜色样式已忽略" in warning for warning in warnings)
    assert any("删除线样式已忽略" in warning for warning in warnings)
    assert any("fbox 已按 Word 原生公式框显示" in warning for warning in warnings)
    assert {node.get(qn("w:val")) for node in document._element.xpath(".//w:color")} >= {
        "FF0000",
        "0000FF",
    }
    assert "FFFF00" in {
        node.get(qn("w:fill")) for node in document._element.xpath(".//w:shd")
    }
    assert sum(r"\begin{CD}" in source for source in rendered) == 2
    assert all(r"\xymatrix" not in source and r"\begin{tikzcd}" not in source for source in rendered)
    assert rendered[0] == r"J05:\quad \ce{2H2 + O2 -> 2H2O}+\pu{9.81 m s^-2}"


def test_browser_renderer_hides_windows_process(monkeypatch):
    captured: dict[str, object] = {}

    def fake_run(command, **options):
        captured.update(options)
        return subprocess.CompletedProcess(command, 0, stdout="", stderr="")

    monkeypatch.setattr(browser_capture.subprocess, "run", fake_run)
    browser_capture._run_browser(
        Path("msedge.exe"),
        ["--headless=new"],
        deadline=time.monotonic() + 1,
    )

    if os.name == "nt":
        assert captured["creationflags"] == subprocess.CREATE_NO_WINDOW
        startupinfo = captured["startupinfo"]
        assert startupinfo.dwFlags & subprocess.STARTF_USESHOWWINDOW
        assert startupinfo.wShowWindow == subprocess.SW_HIDE


def test_code_blocks_are_excluded_from_comment_and_note_preprocessing(tmp_path: Path):
    document = _convert_markdown(
        tmp_path,
        """```html
<!-- CODE_COMMENT_MUST_STAY -->
[^code-footnote]: CODE_FOOTNOTE_DEFINITION_MUST_STAY
[^end:code-endnote]: CODE_ENDNOTE_DEFINITION_MUST_STAY
```

正文[^real]。

[^real]: 真实脚注。
""",
    )

    code_text = document.paragraphs[0].text
    assert "<!-- CODE_COMMENT_MUST_STAY -->" in code_text
    assert "[^code-footnote]: CODE_FOOTNOTE_DEFINITION_MUST_STAY" in code_text
    assert "[^end:code-endnote]: CODE_ENDNOTE_DEFINITION_MUST_STAY" in code_text
    assert len(document._element.xpath(".//w:footnoteReference")) == 1


def test_unclosed_html_preserves_source_and_tail_as_editable_content(tmp_path: Path):
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        """<div data-test="UNCLOSED_HTML_BLOCK">
UNCLOSED_HTML_SOURCE_MUST_BE_PRESERVED

## 后续标题

H01_TAIL_AFTER_UNCLOSED_HTML_MUST_BE_EDITABLE

FINAL_DOCUMENT_TAIL_SENTINEL_20260714
""",
    )

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "UNCLOSED_HTML_SOURCE_MUST_BE_PRESERVED" in text
    assert "H01_TAIL_AFTER_UNCLOSED_HTML_MUST_BE_EDITABLE" in text
    assert "FINAL_DOCUMENT_TAIL_SENTINEL_20260714" in text
    assert len(document.inline_shapes) == 0
    assert any("HTML 块未闭合" in warning for warning in warnings)


def test_html_sanitizer_keeps_resources_while_removing_executable_content():
    safe = browser_renderer.sanitize_html_fragment(
        """<div onclick="alert(1)" style="background:url('https://example.com/bg.png')">
<img src="file:///C:/Windows/Web/Wallpaper/Windows/img0.jpg">
<img src="data:image/png;base64,AAAA">
<script>alert(1)</script>
</div>"""
    )

    assert "file:///C:/Windows/Web/Wallpaper/Windows/img0.jpg" in safe
    assert "https://example.com/bg.png" in safe
    assert "data:image/png;base64,AAAA" in safe
    assert "onclick" not in safe
    assert "<script" not in safe


def test_markdown_images_can_use_absolute_paths_outside_input_directory(tmp_path: Path):
    workspace = tmp_path / "workspace"
    docs = workspace / "docs"
    docs.mkdir(parents=True)
    inside = workspace / "inside.png"
    outside = tmp_path / "outside.png"
    inside.write_bytes(ONE_PIXEL_PNG)
    outside.write_bytes(ONE_PIXEL_PNG)

    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        f"![inside](../inside.png)\n\n![outside]({outside.as_posix()})",
        base_path=docs,
    )

    assert len(document.inline_shapes) == 2
    assert warnings == []


def test_overwrite_failure_preserves_existing_docx(tmp_path: Path, monkeypatch):
    output_path = tmp_path / "existing.docx"
    original = b"ORIGINAL_DOCX_BYTES"
    output_path.write_bytes(original)

    def fail_postprocess(*args, **kwargs):
        raise RuntimeError("injected postprocess failure")

    monkeypatch.setattr(docx_package, "postprocess_docx_package", fail_postprocess)
    document = Document()

    with pytest.raises(RuntimeError, match="injected postprocess failure"):
        docx_package.save_document_atomically(
            document,
            output_path,
            footnotes=[],
            endnotes=[],
            update_fields=False,
            overwrite=True,
        )

    assert output_path.read_bytes() == original
    assert list(tmp_path.glob(".existing-*")) == []


def test_markdown_docx_service_returns_reusable_docx_bytes(tmp_path: Path):
    result = MarkdownDocxService().convert("# 公共转换服务", base_path=tmp_path)

    document = Document(BytesIO(result.content))

    assert document.paragraphs[0].text == "公共转换服务"
    assert result.warnings == ()


def test_markdown_docx_service_forwards_page_layout_options(tmp_path: Path):
    result = MarkdownDocxService().convert(
        "# 横向 A4",
        base_path=tmp_path,
        page_orientation="landscape",
        page_size="a4",
    )

    document = Document(BytesIO(result.content))
    section = document.sections[0]

    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width > section.page_height
    assert round(section.page_width.mm) == 297
    assert round(section.page_height.mm) == 210


def test_unclosed_display_delimiter_does_not_consume_following_content(tmp_path: Path):
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        """\\[

后续内容仍然存在。
""",
    )

    assert warnings == [r"公式定界符未闭合，已保留原文：\["]
    assert [paragraph.text for paragraph in document.paragraphs] == [r"\[", "后续内容仍然存在。"]


def test_latex_validation_rejects_mismatched_structures():
    assert validate_latex(r"\begin{pmatrix}a&b\end{bmatrix}") == (
        "LaTeX 环境开始和结束标签不匹配：bmatrix"
    )
    assert validate_latex(r"\left(x+1") == r"\left 和 \right 未配对"
    assert validate_latex(r"\begin{tikzcd}A\end{tikzcd}") == ""


def test_diagram_normalization_rejects_unsupported_arrow_directions():
    source = r"\xymatrix{A\ar[dr]^f&B\\C&D}"

    try:
        normalize_for_image(source)
    except ValueError as exc:
        assert "暂不支持 dr 方向箭头" in str(exc)
    else:
        raise AssertionError("不支持的交换图方向不应静默生成图片")


def _convert_markdown(
    tmp_path: Path,
    markdown: str,
    *,
    base_path: Path | None = None,
) -> Document:
    document, warnings = _convert_markdown_with_warnings(
        tmp_path,
        markdown,
        base_path=base_path,
    )
    assert warnings == []
    return document


def _convert_markdown_with_warnings(
    tmp_path: Path,
    markdown: str,
    *,
    base_path: Path | None = None,
) -> tuple[Document, list[str]]:
    output_path = tmp_path / "output.docx"
    warnings = convert_markdown_to_docx(
        markdown,
        output_path,
        base_path=base_path or tmp_path,
    )
    return Document(output_path), warnings


def _table_rows(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def _mock_formula_renderer(monkeypatch, rendered: list[str]) -> None:
    def fake_render(
        source: str,
        output_path: Path,
        *,
        display_mode: bool = True,
        budget=None,
    ) -> None:
        rendered.append(source)
        output_path.write_bytes(ONE_PIXEL_PNG)

    monkeypatch.setattr(browser_renderer, "render_katex_png", fake_render)
