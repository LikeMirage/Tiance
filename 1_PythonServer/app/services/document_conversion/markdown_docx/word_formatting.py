from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.image.image import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .word_table_layout import document_available_width_twips
from .word_xml import get_or_add_ordered_child

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMML = f"{{{OMML_NS}}}"
DEFAULT_CHINESE_FONT = "微软雅黑"
DEFAULT_ENGLISH_FONT = "Times New Roman"
DEFAULT_MATH_FONT = "Cambria Math"
EMOJI_FONT = "Segoe UI Emoji"
EMOJI_CLUSTER_RE = re.compile(
    r"(?:[\U0001F1E6-\U0001F1FF]{2}|"
    r"[\U0001F300-\U0001FAFF\u2600-\u27BF][\ufe0f\U0001F3FB-\U0001F3FF]?"
    r"(?:\u200d[\U0001F300-\U0001FAFF\u2600-\u27BF][\ufe0f\U0001F3FB-\U0001F3FF]?)*"
    r")"
)


class FontSettings:
    def __init__(
        self,
        *,
        chinese: str = DEFAULT_CHINESE_FONT,
        english: str = DEFAULT_ENGLISH_FONT,
        math: str = DEFAULT_MATH_FONT,
    ) -> None:
        self.chinese = chinese
        self.english = english
        self.math = math


def block_image_max_width(document):
    return Inches(document_available_width_twips(document) / 1440)


def block_image_max_height(document):
    section = document.sections[-1]
    available = int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)
    return max(Inches(1), int(available * 0.9))


def image_render_width(image_path: Path, max_width):
    try:
        image = Image.from_file(str(image_path))
        if image.width and int(image.width) > 0:
            return min(image.width, max_width)
    except Exception:
        pass
    return max_width


def image_render_dimensions(image_path: Path, max_width, max_height=None):
    try:
        image = Image.from_file(str(image_path))
        width = int(image.width)
        height = int(image.height)
        if width > 0 and height > 0:
            scale = min(1.0, int(max_width) / width)
            if max_height is not None:
                scale = min(scale, int(max_height) / height)
            return max(1, round(width * scale)), max(1, round(height * scale))
    except Exception:
        pass
    return max_width, None


def formula_image_render_width(image_path: Path, max_width):
    high_density_width = image_render_width(image_path, max_width * 2)
    return min(max_width, max(Inches(0.25), high_density_width // 2))


def apply_run_fonts(run, fonts: FontSettings) -> None:
    if _is_emoji_run(run):
        return
    run.font.name = fonts.english
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), fonts.english)
    r_fonts.set(qn("w:hAnsi"), fonts.english)
    r_fonts.set(qn("w:cs"), fonts.english)
    r_fonts.set(qn("w:eastAsia"), fonts.chinese)


def apply_emoji_run_font(run) -> None:
    run.font.name = EMOJI_FONT
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), EMOJI_FONT)
    r_fonts.set(qn("w:hAnsi"), EMOJI_FONT)
    r_fonts.set(qn("w:cs"), EMOJI_FONT)
    r_fonts.set(qn("w:eastAsia"), EMOJI_FONT)


def apply_math_run_font(run, fonts: FontSettings) -> None:
    run.font.name = fonts.math
    r_fonts = run_r_fonts(run)
    r_fonts.set(qn("w:ascii"), fonts.math)
    r_fonts.set(qn("w:hAnsi"), fonts.math)
    r_fonts.set(qn("w:cs"), fonts.math)
    r_fonts.set(qn("w:eastAsia"), fonts.math)


def set_style_fonts(style, fonts: FontSettings) -> None:
    properties = style._element.get_or_add_rPr()
    r_fonts = get_or_add_ordered_child(properties, "w:rFonts")
    r_fonts.set(qn("w:ascii"), fonts.english)
    r_fonts.set(qn("w:hAnsi"), fonts.english)
    r_fonts.set(qn("w:cs"), fonts.english)
    r_fonts.set(qn("w:eastAsia"), fonts.chinese)


def run_r_fonts(run):
    properties = run._element.get_or_add_rPr()
    return get_or_add_ordered_child(properties, "w:rFonts")


def apply_omml_font(root: Any, fonts: FontSettings, size=None) -> None:
    for math_run in root.iter(f"{OMML}r"):
        properties = math_run.find(qn("w:rPr"))
        if properties is None:
            properties = OxmlElement("w:rPr")
            math_properties = math_run.find(f"{OMML}rPr")
            math_run.insert(1 if math_properties is not None else 0, properties)
        r_fonts = get_or_add_ordered_child(properties, "w:rFonts")
        r_fonts.set(qn("w:ascii"), fonts.math)
        r_fonts.set(qn("w:hAnsi"), fonts.math)
        r_fonts.set(qn("w:cs"), fonts.math)
        r_fonts.set(qn("w:eastAsia"), fonts.math)
        if size is not None:
            half_points = str(max(1, round(size.pt * 2)))
            get_or_add_ordered_child(properties, "w:sz").set(qn("w:val"), half_points)
            get_or_add_ordered_child(properties, "w:szCs").set(qn("w:val"), half_points)


def add_plain_runs(paragraph, text: str) -> None:
    position = 0
    for match in EMOJI_CLUSTER_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        run = paragraph.add_run(match.group(0))
        apply_emoji_run_font(run)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def normalize_paragraph(paragraph, *, first_line: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)


def normalize_heading_paragraph(paragraph, level: int) -> None:
    paragraph.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    paragraph.paragraph_format.space_after = Pt(6 if level <= 3 else 4)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = True


def heading_font_size(level: int):
    if level <= 1:
        return Pt(18)
    if level == 2:
        return Pt(15)
    if level == 3:
        return Pt(13)
    return Pt(12)


def normalize_code_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)


def normalize_table_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def style_runs(runs, fonts: FontSettings) -> None:
    for run in runs:
        if not run.font.name:
            apply_run_fonts(run, fonts)
        run.font.size = run.font.size or Pt(12)


def set_runs_default_size(runs, size) -> None:
    for run in runs:
        if run.font.size is None or run.font.size == Pt(12):
            run.font.size = size


def apply_alignment(paragraph, alignment: str) -> None:
    if alignment == "center":
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif alignment == "right":
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def set_shading(paragraph, color: str) -> None:
    shading = get_or_add_ordered_child(paragraph._p.get_or_add_pPr(), "w:shd")
    shading.set(qn("w:fill"), color)


def set_paragraph_border(paragraph, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)


def set_left_border(paragraph, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    border = borders.find(qn("w:left"))
    if border is None:
        border = OxmlElement("w:left")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), color)


def set_horizontal_rule(paragraph, color: str = "B7B7B7") -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = get_or_add_ordered_child(paragraph_properties, "w:pBdr")
    border = borders.find(qn("w:bottom"))
    if border is None:
        border = OxmlElement("w:bottom")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def _is_emoji_run(run) -> bool:
    if run.font.name == EMOJI_FONT:
        return True
    text = run.text or ""
    return bool(text and EMOJI_CLUSTER_RE.fullmatch(text))
