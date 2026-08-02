from __future__ import annotations

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt

from . import markdown_tables
from . import word_formatting as formatting
from . import word_table_layout as table_formatting
from .inline_writer import InlineWriter
from .table_layout import calculate_column_widths
from .text_measurement import FontTextMeasurer
from .word_formatting import FontSettings


TABLE_FONT_SIZE_POINTS = 10.5


class TableWriter:
    """Parses and writes one Markdown pipe table."""

    def __init__(self, document, fonts: FontSettings, inline: InlineWriter) -> None:
        self._document = document
        self._fonts = fonts
        self._inline = inline
        self._measurer = FontTextMeasurer(
            chinese_font=self._fonts.chinese,
            english_font=self._fonts.english,
            math_font=self._fonts.math,
            size_points=TABLE_FONT_SIZE_POINTS,
        )

    def close(self) -> None:
        self._measurer.close()

    def add(self, lines: list[str], start: int) -> int:
        table_lines, end_index = _collect_table_lines(lines, start)
        headers = markdown_tables.parse_table_row(table_lines[0])
        alignments = markdown_tables.parse_table_alignments(table_lines[1])
        rows = [
            markdown_tables.normalize_row(markdown_tables.parse_table_row(line), len(headers))
            for line in table_lines[2:]
        ]
        table = self._document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        available_width_points = table_formatting.document_available_width_points(self._document)
        column_widths = calculate_column_widths(
            headers,
            rows,
            available_width_points=available_width_points,
            cell_padding_points=table_formatting.CELL_HORIZONTAL_PADDING_POINTS,
            measurer=self._measurer,
        )
        table_formatting.apply_column_widths(
            table,
            column_widths,
            self._document,
        )
        table_formatting.set_cell_margins(table)
        table_formatting.set_repeat_header(table.rows[0])
        self._write_header(
            table,
            headers,
            alignments,
            column_widths,
            available_width_points,
        )
        self._write_body(
            table,
            rows,
            alignments,
            column_widths,
            available_width_points,
        )
        return end_index

    def _write_header(
        self,
        table,
        headers: list[str],
        alignments: list[str],
        column_widths: list[float],
        available_width_points: float,
    ) -> None:
        for column, header in enumerate(headers):
            cell = table.rows[0].cells[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_formatting.set_cell_shading(cell, "F2F2F2")
            paragraph = cell.paragraphs[0]
            formatting.normalize_table_paragraph(paragraph)
            self._inline.write(
                paragraph,
                header,
                font_size=Pt(TABLE_FONT_SIZE_POINTS),
                max_image_width=_cell_image_width(
                    column_widths,
                    column,
                    available_width_points,
                ),
            )
            for run in paragraph.runs:
                run.bold = True
            _finish_cell(paragraph, self._fonts)
            formatting.apply_alignment(
                paragraph,
                markdown_tables.header_alignment(alignments, column),
            )

    def _write_body(
        self,
        table,
        rows: list[list[str]],
        alignments: list[str],
        column_widths: list[float],
        available_width_points: float,
    ) -> None:
        for row_index, row_data in enumerate(rows):
            for column, cell_text in enumerate(row_data):
                cell = table.rows[row_index + 1].cells[column]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                formatting.normalize_table_paragraph(paragraph)
                self._inline.write(
                    paragraph,
                    cell_text,
                    font_size=Pt(TABLE_FONT_SIZE_POINTS),
                    max_image_width=_cell_image_width(
                        column_widths,
                        column,
                        available_width_points,
                    ),
                )
                _finish_cell(paragraph, self._fonts)
                formatting.apply_alignment(
                    paragraph,
                    markdown_tables.body_alignment(alignments, column),
                )


def _collect_table_lines(lines: list[str], start: int) -> tuple[list[str], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and "|" in lines[index]:
        table_lines.append(lines[index])
        index += 1
    return table_lines, index - 1


def _finish_cell(paragraph, fonts: FontSettings) -> None:
    formatting.style_runs(paragraph.runs, fonts)
    formatting.set_runs_default_size(paragraph.runs, Pt(10.5))


def _cell_image_width(
    column_widths: list[float],
    column: int,
    available_width_points: float,
):
    percentage = column_widths[column] if column < len(column_widths) else 0.0
    column_width = available_width_points * percentage / 100.0
    content_width = max(18.0, column_width - table_formatting.CELL_HORIZONTAL_PADDING_POINTS)
    return Pt(content_width)
